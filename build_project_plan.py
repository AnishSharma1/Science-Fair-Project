from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION_START

OUT = "MS_EBV_Molecular_Mimicry_Project_Plan.docx"

BLUE = "2E74B5"
DARK = "1F4D78"
PALE = "E8EEF5"
GRAY = "F2F4F7"

def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)

def borders(cell, color="B7C9DB"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:" + edge
        el = tc_borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            tc_borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), color)

def set_cell_margins(cell, top=70, start=120, bottom=70, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn("w:" + side))
        if node is None:
            node = OxmlElement("w:" + side)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")

def font(run, size=10.2, bold=False, color="000000"):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)

def set_style(style, size, color, before, after, bold=False):
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = 1.08

def add_label_para(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(label + " ")
    font(r, 9.5, True, DARK)
    r = p.add_run(text)
    font(r, 9.5)

def main():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.58)
    sec.bottom_margin = Inches(0.58)
    sec.left_margin = Inches(0.7)
    sec.right_margin = Inches(0.7)
    sec.header_distance = Inches(0.3)
    sec.footer_distance = Inches(0.3)
    styles = doc.styles
    set_style(styles["Normal"], 10.2, "000000", 0, 3)
    set_style(styles["Heading 1"], 13.5, BLUE, 8, 3, True)
    set_style(styles["Heading 2"], 10.5, DARK, 5, 2, True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run("RESEARCH EXECUTION BRIEF")
    font(r, 8.5, True, DARK)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("EBV–Myelin Molecular Mimicry in Multiple Sclerosis")
    font(r, 18, True, "0B2545")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run("Goal: build a calibrated, falsifiable pipeline that prioritizes EBV–myelin peptide pairs for TCR cross-reactivity under HLA-DRB1*15:01.")
    font(r, 10, False, "404040")

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(3.43)
    table.columns[1].width = Inches(3.43)
    left, right = table.rows[0].cells
    for c in (left, right):
        shade(c, "F4F6F9"); borders(c); set_cell_margins(c); c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    lp = left.paragraphs[0]; lp.paragraph_format.space_after = Pt(2)
    rr = lp.add_run("Primary hypothesis") ; font(rr, 9.5, True, DARK)
    p = left.add_paragraph(); p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Some EBV and myelin peptides share a sufficiently similar HLA-II-facing / TCR-facing presentation to support plausible cross-recognition."); font(r, 9.2)
    rp = right.paragraphs[0]; rp.paragraph_format.space_after = Pt(2)
    rr = rp.add_run("Claim discipline") ; font(rr, 9.5, True, DARK)
    p = right.add_paragraph(); p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Computational results prioritize candidates; they do not demonstrate molecular mimicry or pathogenicity without functional validation."); font(r, 9.2)

    doc.add_heading("1. Lock the analysis before modeling", level=1)
    add_label_para(doc, "Scope.", "Analyze class II candidates only: HLA-DRB1*15:01, EBV peptides, and myelin / MS-relevant self peptides. Keep peptide lengths and binding checks class-II appropriate.")
    add_label_para(doc, "Candidate rule.", "Predefine a composite rank from sequence similarity, physicochemical similarity, predicted HLA-II presentation, and biological relevance. Do not re-tune scoring after inspecting the top candidates.")
    add_label_para(doc, "Data audit.", "Reconcile allele labels across all inputs—current code includes DRB1*15:02 controls alongside a DRB1*15:01 target—and document every inclusion, exclusion, and source.")

    doc.add_heading("2. Benchmark set and controls", level=1)
    data = [
        ("Positive controls", "≥3 literature-supported EBV–self / TCR–pMHC cross-reactivity cases, including the known EBV DNA polymerase–MBP case if sequence and structural evidence are available."),
        ("Matched negatives", "≥10–20 pairs matched on allele, peptide length, source class, and binding plausibility but not reported to cross-react. Include shuffled or decoy pairings."),
        ("Exploratory set", "Novel EBV–myelin pairs held out from threshold selection; report all ranks, not only successes."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER; table.autofit = False
    table.columns[0].width = Inches(1.32); table.columns[1].width = Inches(5.54)
    for i, hdr in enumerate(["Set", "Purpose / acceptance criterion"]):
        c = table.rows[0].cells[i]; shade(c, PALE); borders(c); set_cell_margins(c)
        r = c.paragraphs[0].add_run(hdr); font(r, 9.3, True, DARK)
    for a, b in data:
        cells = table.add_row().cells
        for c in cells: borders(c); set_cell_margins(c)
        r = cells[0].paragraphs[0].add_run(a); font(r, 9, True, DARK)
        r = cells[1].paragraphs[0].add_run(b); font(r, 9)

    doc.add_heading("3. Modeling and evaluation gates", level=1)
    add_label_para(doc, "Calibration gate.", "A modeling approach advances only if it recovers known positive-control complex geometry and ranks positives above matched negatives. Record RMSD where an experimental structure exists; supplement it with interface confidence and visual interface inspection.")
    add_label_para(doc, "Ranking gate.", "Evaluate enrichment of positives in the top-ranked candidates (e.g., top-k recall, AUROC / AUPRC where labels support it), effect sizes, confidence intervals, and performance versus random and a sequence-only baseline.")
    add_label_para(doc, "Robustness gate.", "Repeat using blinded / held-out controls and sensitivity analyses for peptide window, scoring weights, and candidate-source choice. A single favorable structure is hypothesis-generating only.")

    doc.add_heading("4. Decision path", level=1)
    steps = [
        ("A. Dataset freeze", "Finish provenance table, allele/peptide audit, controls, and predefined metrics."),
        ("B. Method benchmark", "Run AlphaFold 3 and TCRmodel2 on the positive-control set; compare only after the same preprocessing and evaluation protocol."),
        ("C. Candidate triage", "Advance methods that clear calibration; nominate 1–3 novel pairs with confidence, rank, and failure-mode notes."),
        ("D. Validation plan", "For strongest pair(s), propose an HLA-II peptide/T-cell functional assay through an experimental collaborator; RNA-seq remains contextual evidence, not proof of cross-reactivity."),
    ]
    for title, text in steps:
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.left_indent = Inches(0.12); p.paragraph_format.first_line_indent = Inches(-0.12)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(title + " — "); font(r, 9.6, True, BLUE)
        r = p.add_run(text); font(r, 9.6)

    doc.add_heading("First next action", level=1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Build the benchmark manifest before submitting any new structures: exact TCR α/β chains, peptide sequence/window, HLA chains, experimental PDB reference, positive/negative label, and source citation.")
    font(r, 10, True, "0B2545")

    doc.save(OUT)

if __name__ == "__main__":
    main()
